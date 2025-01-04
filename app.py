from flask import Flask, render_template_string, request, redirect, url_for, Response
from googleapiclient.discovery import build
import io
from docx import Document

# Flask uygulaması
app = Flask(__name__)

# YouTube Data API yapılandırması
API_KEY = "AIzaSyBIyDYaA4ZXBKJPl2TktKlWK0een4vnY-8"  # Google API Key

def get_video_links(api_key, playlist_id):
    youtube = build('youtube', 'v3', developerKey=api_key)
    videos = []
    next_page_token = None

    while True:
        request = youtube.playlistItems().list(
            part="snippet",
            playlistId=playlist_id,
            maxResults=50,
            pageToken=next_page_token
        )
        response = request.execute()

        for item in response['items']:
            video_id = item['snippet']['resourceId']['videoId']
            video_url = f"https://www.youtube.com/watch?v={video_id}"
            videos.append(video_url)

        next_page_token = response.get('nextPageToken')
        if not next_page_token:
            break

    return videos

@app.route('/', methods=['GET', 'POST'])
def homepage():
    if request.method == 'POST':
        playlist_url = request.form['playlist_url']
        # Oynatma listesi ID'sini URL'den çıkarıyoruz
        playlist_id = playlist_url.split('list=')[-1]
        return redirect(url_for('display_videos', playlist_id=playlist_id))
    
    html_template = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>YouTube Playlist Extractor</title>
        <style>
            body {
                font-family: Arial, sans-serif;
                background-color: #f4f4f9;
                color: #333;
                margin: 0;
                padding: 0;
                display: flex;
                justify-content: center;
                align-items: center;
                height: 100vh;
            }
            .container {
                text-align: center;
                background: #ffffff;
                padding: 20px;
                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                border-radius: 10px;
            }
            h1 {
                color: #0073e6;
            }
            form {
                margin-top: 20px;
            }
            input[type="text"] {
                padding: 10px;
                width: 80%;
                margin-bottom: 20px;
                border: 1px solid #ddd;
                border-radius: 5px;
                font-size: 16px;
            }
            button {
                padding: 10px 20px;
                background-color: #0073e6;
                color: white;
                border: none;
                border-radius: 5px;
                cursor: pointer;
                font-size: 16px;
            }
            button:hover {
                background-color: #005bb5;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>Enter YouTube Playlist URL</h1>
            <form method="post">
                <input type="text" id="playlist_url" name="playlist_url" placeholder="Enter playlist URL" required><br>
                <button type="submit">Get Video Links</button>
            </form>
        </div>
    </body>
    </html>
    """
    return render_template_string(html_template)

@app.route('/videos/<playlist_id>')
def display_videos(playlist_id):
    video_links = get_video_links(API_KEY, playlist_id)
    html_template = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Video Links</title>
        <style>
            body {
                font-family: Arial, sans-serif;
                background-color: #f4f4f9;
                color: #333;
                margin: 0;
                padding: 0;
                display: flex;
                justify-content: center;
                align-items: center;
                min-height: 100vh;
            }
            .container {
                text-align: center;
                background: #ffffff;
                padding: 20px;
                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                border-radius: 10px;
                width: 90%;
                max-width: 800px;
            }
            h1 {
                color: #0073e6;
            }
            ul {
                list-style: none;
                padding: 0;
            }
            li {
                margin: 10px 0;
            }
            a {
                text-decoration: none;
                color: #0073e6;
            }
            a:hover {
                text-decoration: underline;
            }
            .back-link, .download-link {
                display: block;
                margin-top: 20px;
                color: #0073e6;
                text-decoration: none;
                font-weight: bold;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>YouTube Playlist Video Links</h1>
            <ul>
                {% for link in video_links %}
                <li><a href="{{ link }}" target="_blank">{{ link }}</a></li>
                {% endfor %}
            </ul>
            <a href="{{ url_for('download_txt', playlist_id=playlist_id) }}" class="download-link">Download as Text File</a>
            <a href="{{ url_for('download_docx', playlist_id=playlist_id) }}" class="download-link">Download as Word Document</a>
            <a href="/" class="back-link">Back to Homepage</a>
        </div>
    </body>
    </html>
    """
    return render_template_string(html_template, video_links=video_links, playlist_id=playlist_id)

@app.route('/download_txt/<playlist_id>')
def download_txt(playlist_id):
    video_links = get_video_links(API_KEY, playlist_id)
    text_content = "\n".join(video_links)
    return Response(text_content, mimetype='text/plain', headers={"Content-Disposition": "attachment;filename=video_links.txt"})

@app.route('/download_docx/<playlist_id>')
def download_docx(playlist_id):
    video_links = get_video_links(API_KEY, playlist_id)
    doc = Document()
    doc.add_heading('YouTube Playlist Video Links', level=1)
    for link in video_links:
        doc.add_paragraph(link)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={"Content-Disposition": "attachment;filename=video_links.docx"})

if __name__ == '__main__':
    app.run(debug=True)
